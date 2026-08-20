"""1.2 用户任务文档结构解析 —— Word(.docx) 实现。

职责：把待填 Word 文件解析成「结构树 + 填写位置(locator)」。
- 标题(headings)：带层级
- 段落(paragraphs)：正文
- 表格(tables)：每个单元格内容 + 合并单元格信息
- 可填写位置(fields)：空单元格 + WordCellLocator，供统筹回填

不做：字段业务语义匹配(1.3)、写回文件(统筹)。
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table

from aff_contracts import FormStructureResult

from app.modules.form_structure.port import FormStructurePort


# Word 中标题样式的常见前缀（中英文）
_HEADING_STYLE_PREFIXES = ("Heading", "标题", "Title")

# 填空模板正则：日期、地址、盖章等固定提示文字
_DATE_TEMPLATE = re.compile(r"年\s+月\s+日")
_ADDR_TEMPLATE = re.compile(r"省\s+市\s+[县区]")
_SEAL_PATTERN = re.compile(r"盖章|签字|签名|签收")
# 格式说明词（短词，表示填写格式要求，非数据也非标签）
_FORMAT_HINT_WORDS = {"全称", "简称", "原名", "中文名", "英文名", "中文", "英文"}
# 带括号且括号内含数字的格式示例（如 2/50（4%）、3/100（5%））
_FORMAT_EXAMPLE_PATTERN = re.compile(r"[\(（]\s*\d")


def _is_format_hint(text: str) -> bool:
    """检测格式提示文字（告诉用户怎么填的简短说明，本身不是数据）。

    例：「全称」（要求填全称）、「2/50（4%）」（格式示例）。
    区别于预填数据：「4课时」是已填好的值，不含括号也不是格式词。
    """
    if text in _FORMAT_HINT_WORDS:
        return True
    if _FORMAT_EXAMPLE_PATTERN.search(text):
        return True
    return False


def _is_instruction_text(text: str) -> bool:
    """检测填写说明文字（有文字但是填写区，不是纯标签）。

    纯标签如「姓名」「项目名称」不含句子标点和指令词；
    填写说明如「大学期间获得的校级及以上奖励，如果无校级及以上，则院级也可以」
    「时间段＋具体岗位」含指令词/标点/填写提示符。
    """
    # 含填写提示符（全角/半角加号表示并列填写项）
    if re.search(r"[＋+]", text):
        return True
    # 含条件/建议指令词
    if any(kw in text for kw in ("如果", "如无", "请填写", "请写", "请如实")):
        return True
    # 含中文句子标点 + 长度>10 → 说明书/提示语
    if len(text) > 10 and re.search(r"[，。；]", text):
        return True
    return False


def _is_fillable_template(text: str) -> bool:
    """检测单元格是否是带固定提示文字的填空栏（非空但需要填写）。

    例：「年  月  日」「省  市  县  镇  村」「盖章：\\n年  月  日」
    「辅导员签字：」「大学期间获得的校级及以上奖励，...」
    这些单元格有文字但本质是填写区，不能当纯标签跳过。
    """
    if not text:
        return False
    if _DATE_TEMPLATE.search(text):
        return True
    if _ADDR_TEMPLATE.search(text):
        return True
    if _SEAL_PATTERN.search(text):
        return True
    if re.search(r"_{2,}", text):  # 下划线占位
        return True
    if _is_instruction_text(text):
        return True
    if _is_format_hint(text):
        return True
    return False


def _is_heading(style_name: str) -> bool:
    """判断段落样式是否是标题。"""
    return any(style_name.startswith(p) for p in _HEADING_STYLE_PREFIXES)


def _heading_level(style_name: str) -> int:
    """从样式名提取标题层级，如 'Heading 2' / '标题 1' → 2 / 1。取不到数字时按 1 级。"""
    for token in style_name.split():
        if token.isdigit():
            return int(token)
    return 1


def _extract_headings(doc: DocxDocument) -> list[dict[str, Any]]:
    """提取所有标题段落。"""
    headings: list[dict[str, Any]] = []
    for idx, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if _is_heading(style):
            headings.append(
                {
                    "index": idx,
                    "level": _heading_level(style),
                    "text": text,
                    "style": style,
                }
            )
    return headings


def _extract_paragraphs(doc: DocxDocument) -> list[dict[str, Any]]:
    """提取所有非空正文段落（含标题，按原文顺序，便于定位）。"""
    paragraphs: list[dict[str, Any]] = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else "Normal"
        paragraphs.append({"index": idx, "text": text, "style": style})
    return paragraphs


def _build_tc_grid(table: Table, num_rows: int, num_cols: int) -> list[list[Any]]:
    """预读所有 cell 的 _tc 引用并保持存活，防止 GC 导致 id() 地址复用。

    python-docx 的 table.cell(r,c) 每次返回新的 _Cell 包装对象；
    若不保持引用，旧对象被回收后 id() 会被新对象复用，造成「假合并」。
    合并单元格的多个 (r,c) 会指向同一个底层 <w:tc> 元素（_tc 相同）。
    """
    grid: list[list[Any]] = []
    for r in range(num_rows):
        row: list[Any] = []
        for c in range(num_cols):
            cell = table.cell(r, c)
            # 同时保持 cell 和 _tc 引用存活
            row.append({"cell": cell, "tc": cell._tc, "text": cell.text.strip()})
        grid.append(row)
    return grid


def _detect_header_rows(grid: list[list[dict[str, Any]]], num_rows: int, num_cols: int) -> int:
    """启发式推断表头行数（含多级表头）。

    规则：
    - 找到第一个含空单元格的行 → 该行之前都是表头（返回该行号）
    - form 式标签同行表（标签与空格同行）→ 返回 0（无独立表头行）
    - 无空单元格的展示表 → 返回 1（约定首行为表头）
    - 空表 → 返回 0

    配合 merged_ranges 可还原多级表头：header_rows > 1 且首行有 colspan 合并
    即为多级表头。
    """
    if num_rows == 0:
        return 0
    for r in range(num_rows):
        for c in range(num_cols):
            text = grid[r][c]["text"]
            if not text or _is_fillable_template(text):
                return r  # 第 r 行有可填写格，之前都是表头
    return 1  # 无可填写格，约定首行为表头


def _extract_table(table: Table, table_index: int) -> dict[str, Any]:
    """解析单个表格：单元格内容 + 合并单元格范围 + 表头行数。

    依赖 _build_tc_grid 保持所有 _tc 引用存活，此时 id(_tc) 稳定可靠：
    合并单元格的多个 (r,c) 拥有相同 id(_tc)，非合并的各不相同。
    """
    num_rows = len(table.rows)
    num_cols = len(table.columns) if num_rows else 0
    grid = _build_tc_grid(table, num_rows, num_cols)
    header_rows = _detect_header_rows(grid, num_rows, num_cols)

    cells: list[dict[str, Any]] = []
    # id(tc) → 该合并块的边界信息
    ranges: dict[int, dict[str, Any]] = {}

    for r in range(num_rows):
        for c in range(num_cols):
            info = grid[r][c]
            tc_id = id(info["tc"])
            text = info["text"]
            if tc_id not in ranges:
                # 首次见到这个 _tc（合并块的左上角）
                ranges[tc_id] = {
                    "min_row": r,
                    "min_col": c,
                    "max_row": r,
                    "max_col": c,
                }
                cells.append(
                    {
                        "row": r,
                        "col": c,
                        "text": text,
                        "merged": False,  # 是否被合并（>1 格）
                        "_tc_id": tc_id,  # 内部字段，输出前清理
                    }
                )
            else:
                # 同一个 _tc 的其它位置：扩展合并范围
                ranges[tc_id]["max_row"] = max(ranges[tc_id]["max_row"], r)
                ranges[tc_id]["max_col"] = max(ranges[tc_id]["max_col"], c)

    # 标记被合并的单元格，整理合并范围
    merged_ranges: list[dict[str, Any]] = []
    for cell_info in cells:
        rng = ranges[cell_info["_tc_id"]]
        span_rows = rng["max_row"] - rng["min_row"] + 1
        span_cols = rng["max_col"] - rng["min_col"] + 1
        if span_rows > 1 or span_cols > 1:
            cell_info["merged"] = True
            cell_info["rowspan"] = span_rows
            cell_info["colspan"] = span_cols
            merged_ranges.append(rng)
        cell_info.pop("_tc_id", None)  # 清理内部字段，不进输出

    return {
        "index": table_index,
        "rows": num_rows,
        "cols": num_cols,
        "cells": cells,
        "merged_ranges": merged_ranges,
        "header_rows": header_rows,
    }


def _clean_label(text: str) -> str:
    """清理标签文字：去除换行和首尾空白。"""
    return text.replace("\n", "").replace("\r", "").strip()


def _is_vertically_merged(
    grid: list[list[dict[str, Any]]], row: int, col: int
) -> bool:
    """检测单元格是否是纵向合并的延续部分（即上方行同列是同一个 <w:tc>）。

    纵向合并的栏目标题（如「家庭成员情况」跨行 4-10）会在每行都返回相同文字，
    但它本质是「小节标题」而非「逐行标签」，扫描时应跳过。
    """
    return row > 0 and grid[row][col]["tc"] is grid[row - 1][col]["tc"]


def _find_hmerge_origin_text(
    grid: list[list[dict[str, Any]]], row: int, col: int
) -> str:
    """如果 (row, col) 是横向合并的延续格（text 为空），向左找到源头格的文字。

    例：行4 col10-12 合并为「年收入（元）」，col11 的 text 为空，
    本函数向左扫描找到 col10 的「年收入（元）」。
    """
    if grid[row][col]["text"]:
        return ""
    for c in range(col - 1, -1, -1):
        if grid[row][c]["tc"] is grid[row][col]["tc"] and grid[row][c]["text"]:
            return grid[row][c]["text"]
    return ""


def _scan_for_label(
    grid: list[list[dict[str, Any]]], row: int, col: int, num_cols: int
) -> str:
    """向上 + 向左扫描最近的有意义标签文字，作为字段名。

    扫描顺序：
    1. 同行向左（跳过空格、模板文字、纵向合并的小节标题）—— 适用于「标签 | 填写格」同行布局
    2. 同列向上（跳过空格和模板文字；空格时查找横合并源头）—— 适用于表头在上方多行处
    3. 同行 col0（跳过纵向合并）—— 适用于「意见」类栏目
    4. 向上扫 col0（跳过纵向合并）—— 兜底
    """
    # 1. 同行向左（跳过纵向合并的小节标题）
    for c in range(col - 1, -1, -1):
        text = grid[row][c]["text"]
        if not text or _is_fillable_template(text):
            continue
        if _is_vertically_merged(grid, row, c):
            continue  # 纵向合并 → 小节标题，跳过
        return _clean_label(text)
    # 2. 同列向上（空格时查找横合并源头）
    for r in range(row - 1, -1, -1):
        text = grid[r][col]["text"]
        if text and not _is_fillable_template(text):
            return _clean_label(text)
        # 空格 → 可能是横合并的延续格，向左找源头
        hmerge_text = _find_hmerge_origin_text(grid, r, col)
        if hmerge_text and not _is_fillable_template(hmerge_text):
            return _clean_label(hmerge_text)
    # 3. 同行 col0（跳过纵向合并的小节标题）
    if col > 0:
        text = grid[row][0]["text"]
        if text and not _is_fillable_template(text):
            if not _is_vertically_merged(grid, row, 0):
                return _clean_label(text)
    # 4. 向上扫 col0（跳过纵向合并）
    for r in range(row - 1, -1, -1):
        text = grid[r][0]["text"]
        if text and not _is_fillable_template(text):
            if not _is_vertically_merged(grid, r, 0):
                return _clean_label(text)
    return ""


def _extract_fields_from_tables(
    tables: list[Table],
) -> list[dict[str, Any]]:
    """识别可填写位置：空单元格 + 模板填空栏 → WordCellLocator。

    可填写判定：
    - 单元格为空 → 可填写
    - 单元格含模板提示文字（日期/地址/盖章等）→ 可填写
    - 单元格是纯标签文字 → 不可填写（跳过）

    字段名用 _scan_for_label 向上向左扫描最近标签，消除(未知)。
    """
    fields: list[dict[str, Any]] = []
    for t_idx, table in enumerate(tables):
        num_rows = len(table.rows)
        num_cols = len(table.columns) if num_rows else 0
        grid = _build_tc_grid(table, num_rows, num_cols)  # 保持 _tc 引用存活
        seen_tcs: set[int] = set()  # 已处理过的 _tc 元素 id
        for r in range(num_rows):
            for c in range(num_cols):
                info = grid[r][c]
                tc_id = id(info["tc"])
                # 只在每个合并块的左上角处理一次
                if tc_id in seen_tcs:
                    continue
                seen_tcs.add(tc_id)
                text = info["text"]
                # 纵向合并的延续格（非左上角）→ 不是独立可填写位，跳过
                if _is_vertically_merged(grid, r, c):
                    continue
                # 空格 → 可填写；模板文字 → 可填写；纯标签 → 跳过
                if text and not _is_fillable_template(text):
                    continue  # 纯标签，不是可填写位
                name = _scan_for_label(grid, r, c, num_cols)
                fill_kind = "template" if text else "empty"
                fields.append(
                    {
                        "id": f"word_t{t_idx}_r{r}_c{c}",
                        "name": name,
                        "fill_kind": fill_kind,
                        "locator": {
                            "kind": "word_cell",
                            "table_index": t_idx,
                            "row": r,
                            "col": c,
                        },
                    }
                )
    return fields


class FormStructureService(FormStructurePort):
    """1.2 Word 文档结构解析服务。"""

    def parse_structure(self, file_path: str, *, filename: str) -> dict[str, Any]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        if ext not in {".docx", ".doc"}:
            raise ValueError(f"1.2 第一阶段仅支持 Word(.docx)，收到: {ext}")
        if ext == ".doc":
            raise ValueError(
                "旧版 .doc 格式不支持，请另存为 .docx 后再上传"
            )

        doc = Document(str(path))

        headings = _extract_headings(doc)
        paragraphs = _extract_paragraphs(doc)
        tables_raw = list(doc.tables)
        tables = [_extract_table(t, i) for i, t in enumerate(tables_raw)]
        fields = _extract_fields_from_tables(tables_raw)

        notes: list[str] = []
        if not headings:
            notes.append("未检测到标题段落（可能使用了自定义样式）")
        if not tables_raw:
            notes.append("文档无表格，无可填写位置")
        else:
            notes.append(f"共解析 {len(tables_raw)} 个表格，识别 {len(fields)} 个可填写位置")

        result = {
            "format": "docx",
            "filename": filename,
            "title": headings[0]["text"] if headings else path.stem,
            "headings": headings,
            "paragraphs": paragraphs,
            "tables": tables,
            "fields": fields,
            "structure_notes": notes,
        }
        # 契约校验：实现与 FormStructureResult 契约漂移时立即报错
        return FormStructureResult(**result).model_dump()


def get_form_structure_service() -> FormStructurePort:
    return FormStructureService()
