from pathlib import Path

from aff_contracts import FileRef, IngestRequest
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter
from pypdf.generic import (
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)

from app.modules.ingest.service import IngestService


def make_request(
    path: Path,
    mime: str,
) -> IngestRequest:
    return IngestRequest(
        doc_id="doc_test_001",
        file=FileRef(
            path=str(path),
            filename=path.name,
            mime=mime,
            sha256="test",
            size=path.stat().st_size,
        ),
    )


def test_ingest_txt(tmp_path: Path) -> None:
    path = tmp_path / "sample.txt"

    path.write_text(
        "这是 TXT 测试。",
        encoding="utf-8",
    )

    request = make_request(
        path,
        "text/plain",
    )

    bundle = IngestService().ingest(request)

    assert bundle.text == "这是 TXT 测试。"
    assert bundle.title == "sample"
    assert bundle.media_type == "text/plain"


def test_ingest_markdown(
    tmp_path: Path,
) -> None:
    path = tmp_path / "sample.md"

    path.write_text(
        "# 个人经历\n\n这是 Markdown 测试。",
        encoding="utf-8",
    )

    request = make_request(
        path,
        "text/markdown",
    )

    bundle = IngestService().ingest(request)

    assert bundle.text == (
        "# 个人经历\n\n这是 Markdown 测试。"
    )

    assert bundle.title == "sample"
    assert bundle.media_type == "text/markdown"


def test_ingest_docx(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"

    document = Document()

    document.add_heading(
        "个人经历",
        level=1,
    )

    document.add_paragraph(
        "赵洋帆就读于湖南大学。"
    )

    document.add_paragraph(
        "负责知识图谱相关项目研究。"
    )

    document.save(path)

    request = make_request(
        path,
        (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )

    bundle = IngestService().ingest(request)

    assert bundle.text == (
        "个人经历\n"
        "赵洋帆就读于湖南大学。\n"
        "负责知识图谱相关项目研究。"
    )

    assert bundle.title == "sample"


def test_ingest_xlsx(tmp_path: Path) -> None:
    path = tmp_path / "sample.xlsx"

    workbook = Workbook()
    sheet = workbook.active

    sheet.title = "基本信息"

    sheet.append(
        [
            "姓名",
            "学校",
            "专业",
        ]
    )

    sheet.append(
        [
            "赵洋帆",
            "湖南大学",
            "大数据管理与应用",
        ]
    )

    workbook.save(path)
    workbook.close()

    request = make_request(
        path,
        (
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
    )

    bundle = IngestService().ingest(request)

    assert bundle.text == (
        "[Sheet: 基本信息]\n"
        "姓名\t学校\t专业\n"
        "赵洋帆\t湖南大学\t大数据管理与应用"
    )

    assert bundle.title == "sample"

    assert bundle.metadata["sheet_names"] == [
        "基本信息",
    ]


def test_ingest_pdf(tmp_path: Path) -> None:
    path = tmp_path / "sample.pdf"

    writer = PdfWriter()

    page = writer.add_blank_page(
        width=612,
        height=792,
    )

    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject(
                "/Type1"
            ),
            NameObject("/BaseFont"): NameObject(
                "/Helvetica"
            ),
        }
    )

    font_ref = writer._add_object(font)

    page[NameObject("/Resources")] = (
        DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {
                        NameObject("/F1"): font_ref,
                    }
                )
            }
        )
    )

    stream = DecodedStreamObject()

    stream.set_data(
        b"BT /F1 12 Tf 72 720 Td "
        b"(Hello PDF) Tj ET"
    )

    stream_ref = writer._add_object(stream)

    page[NameObject("/Contents")] = stream_ref

    writer.write(path)

    request = make_request(
        path,
        "application/pdf",
    )

    bundle = IngestService().ingest(request)

    assert bundle.text == "Hello PDF"
    assert bundle.title == "sample"
    assert bundle.media_type == "application/pdf"

    assert bundle.metadata["page_count"] == 1


def test_ingest_xlsx_tables(
    tmp_path: Path,
) -> None:
    path = tmp_path / "awards.xlsx"

    workbook = Workbook()
    sheet = workbook.active

    sheet.title = "获奖情况"

    sheet.append(
        [
            "获奖时间",
            "奖项名称",
            "级别",
        ]
    )

    sheet.append(
        [
            "2023-11",
            "国家奖学金",
            "国家级",
        ]
    )

    sheet.append(
        [
            "2024-05",
            "挑战杯",
            "省级金奖",
        ]
    )

    workbook.save(path)
    workbook.close()

    request = make_request(
        path,
        (
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
    )

    bundle = IngestService().ingest(request)

    assert bundle.tables is not None
    assert len(bundle.tables) == 1

    table = bundle.tables[0]

    assert table.name == "获奖情况"

    assert table.headers == [
        "获奖时间",
        "奖项名称",
        "级别",
    ]

    assert table.rows == [
        [
            "2023-11",
            "国家奖学金",
            "国家级",
        ],
        [
            "2024-05",
            "挑战杯",
            "省级金奖",
        ],
    ]


def test_ingest_docx_tables(
    tmp_path: Path,
) -> None:
    path = tmp_path / "profile.docx"

    document = Document()

    document.add_heading(
        "个人信息",
        level=1,
    )

    table = document.add_table(
        rows=3,
        cols=3,
    )

    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "学校"
    table.cell(0, 2).text = "专业"

    table.cell(1, 0).text = "赵洋帆"
    table.cell(1, 1).text = "湖南大学"
    table.cell(1, 2).text = "大数据管理与应用"

    table.cell(2, 0).text = "张三"
    table.cell(2, 1).text = "湖南大学"
    table.cell(2, 2).text = "计算机科学与技术"

    document.save(path)

    request = make_request(
        path,
        (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )

    bundle = IngestService().ingest(request)

    assert bundle.tables is not None
    assert len(bundle.tables) == 1

    table_result = bundle.tables[0]

    assert table_result.name == "Table 1"

    assert table_result.headers == [
        "姓名",
        "学校",
        "专业",
    ]

    assert table_result.rows == [
        [
            "赵洋帆",
            "湖南大学",
            "大数据管理与应用",
        ],
        [
            "张三",
            "湖南大学",
            "计算机科学与技术",
        ],
    ]

    assert "个人信息" in bundle.text
    assert "赵洋帆" in bundle.text
    assert "湖南大学" in bundle.text


def test_ingest_pdf_chunks(
    tmp_path: Path,
) -> None:
    path = tmp_path / "chunks.pdf"

    writer = PdfWriter()

    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject(
                "/Type1"
            ),
            NameObject("/BaseFont"): NameObject(
                "/Helvetica"
            ),
        }
    )

    font_ref = writer._add_object(font)

    texts = [
        b"First Page",
        b"Second Page",
    ]

    for text_value in texts:
        page = writer.add_blank_page(
            width=612,
            height=792,
        )

        page[NameObject("/Resources")] = (
            DictionaryObject(
                {
                    NameObject(
                        "/Font"
                    ): DictionaryObject(
                        {
                            NameObject(
                                "/F1"
                            ): font_ref,
                        }
                    )
                }
            )
        )

        stream = DecodedStreamObject()

        stream.set_data(
            b"BT /F1 12 Tf 72 720 Td ("
            + text_value
            + b") Tj ET"
        )

        stream_ref = writer._add_object(
            stream
        )

        page[
            NameObject("/Contents")
        ] = stream_ref

    writer.write(path)

    request = make_request(
        path,
        "application/pdf",
    )

    bundle = IngestService().ingest(request)

    assert bundle.chunks_hint is not None
    assert len(bundle.chunks_hint) == 2

    assert (
        bundle.chunks_hint[0].text
        == "First Page"
    )
    assert bundle.chunks_hint[0].page == 1

    assert (
        bundle.chunks_hint[1].text
        == "Second Page"
    )
    assert bundle.chunks_hint[1].page == 2

    assert bundle.text == (
        "First Page\n\nSecond Page"
    )

    assert bundle.metadata["page_count"] == 2


def test_ingest_xlsx_chunks(
    tmp_path: Path,
) -> None:
    path = tmp_path / "multi_sheet.xlsx"

    workbook = Workbook()

    sheet1 = workbook.active
    sheet1.title = "基本信息"

    sheet1.append(
        [
            "姓名",
            "学校",
            "专业",
        ]
    )

    sheet1.append(
        [
            "赵洋帆",
            "湖南大学",
            "大数据管理与应用",
        ]
    )

    sheet2 = workbook.create_sheet(
        title="获奖情况"
    )

    sheet2.append(
        [
            "获奖时间",
            "奖项名称",
            "级别",
        ]
    )

    sheet2.append(
        [
            "2023-11",
            "国家奖学金",
            "国家级",
        ]
    )

    workbook.save(path)
    workbook.close()

    request = make_request(
        path,
        (
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
    )

    bundle = IngestService().ingest(request)

    assert bundle.chunks_hint is not None
    assert len(bundle.chunks_hint) == 2

    assert (
        bundle.chunks_hint[0].sheet
        == "基本信息"
    )

    assert bundle.chunks_hint[0].text == (
        "姓名\t学校\t专业\n"
        "赵洋帆\t湖南大学\t大数据管理与应用"
    )

    assert (
        bundle.chunks_hint[1].sheet
        == "获奖情况"
    )

    assert bundle.chunks_hint[1].text == (
        "获奖时间\t奖项名称\t级别\n"
        "2023-11\t国家奖学金\t国家级"
    )

    assert bundle.metadata["sheet_names"] == [
        "基本信息",
        "获奖情况",
    ]


def test_ingest_pdf_blank_page_warning(
    tmp_path: Path,
) -> None:
    path = tmp_path / "pdf_with_blank_page.pdf"

    writer = PdfWriter()

    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject(
                "/Type1"
            ),
            NameObject("/BaseFont"): NameObject(
                "/Helvetica"
            ),
        }
    )

    font_ref = writer._add_object(font)

    page1 = writer.add_blank_page(
        width=612,
        height=792,
    )

    page1[NameObject("/Resources")] = (
        DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {
                        NameObject("/F1"): font_ref,
                    }
                )
            }
        )
    )

    stream = DecodedStreamObject()

    stream.set_data(
        b"BT /F1 12 Tf 72 720 Td "
        b"(First Page) Tj ET"
    )

    stream_ref = writer._add_object(stream)

    page1[NameObject("/Contents")] = stream_ref

    writer.add_blank_page(
        width=612,
        height=792,
    )

    writer.write(path)

    request = make_request(
        path,
        "application/pdf",
    )

    bundle = IngestService().ingest(request)

    assert bundle.text == "First Page"

    assert bundle.metadata["page_count"] == 2

    assert bundle.chunks_hint is not None
    assert len(bundle.chunks_hint) == 1

    assert bundle.chunks_hint[0].page == 1

    assert bundle.warnings == [
        (
            "PDF page 2 contains no "
            "extractable text."
        )
    ]


def test_ingest_xlsx_empty_sheet_warning(
    tmp_path: Path,
) -> None:
    path = tmp_path / "excel_with_empty_sheet.xlsx"

    workbook = Workbook()

    sheet1 = workbook.active
    sheet1.title = "基本信息"

    sheet1.append(
        [
            "姓名",
            "学校",
        ]
    )

    sheet1.append(
        [
            "赵洋帆",
            "湖南大学",
        ]
    )

    workbook.create_sheet(
        title="空白表"
    )

    workbook.save(path)
    workbook.close()

    request = make_request(
        path,
        (
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
    )

    bundle = IngestService().ingest(request)

    assert bundle.metadata["sheet_names"] == [
        "基本信息",
        "空白表",
    ]

    assert bundle.tables is not None
    assert len(bundle.tables) == 1

    assert bundle.chunks_hint is not None
    assert len(bundle.chunks_hint) == 1

    assert (
        bundle.chunks_hint[0].sheet
        == "基本信息"
    )

    assert bundle.warnings == [
        "Excel sheet '空白表' is empty."
    ]


def test_ingest_docx_preserves_block_order(
    tmp_path: Path,
) -> None:
    path = tmp_path / "ordered.docx"

    document = Document()

    document.add_heading(
        "一、基本情况",
        level=1,
    )

    document.add_paragraph(
        "下面是基本信息。"
    )

    table = document.add_table(
        rows=2,
        cols=2,
    )

    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "学校"

    table.cell(1, 0).text = "张三"
    table.cell(1, 1).text = "湖南大学"

    document.add_heading(
        "二、补充说明",
        level=1,
    )

    document.add_paragraph(
        "这是表格后面的说明。"
    )

    document.save(path)

    request = make_request(
        path,
        (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )

    bundle = IngestService().ingest(request)

    assert bundle.text == (
        "一、基本情况\n"
        "下面是基本信息。\n"
        "[Table: Table 1]\n"
        "姓名\t学校\n"
        "张三\t湖南大学\n"
        "二、补充说明\n"
        "这是表格后面的说明。"
    )

    assert bundle.tables is not None
    assert len(bundle.tables) == 1
