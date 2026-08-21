"""1.2 form_structure Word 解析测试。

用 python-docx 动态生成样例 docx（含标题/段落/表格/合并单元格/空格），
再验证 parse_structure 的输出。
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from aff_contracts.fill import WordCellLocator
from app.modules.form_structure.service import FormStructureService


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """生成一个样例 Word 文档：
    - 标题：个人信息登记表（Heading 1）/ 基本信息（Heading 2）
    - 段落：一句说明
    - 表格 1：标签+空格的常见填报表（姓名/性别两列，含一个合并单元格）
    - 表格 2：纯展示表（无空格）
    """
    doc = Document()

    # 标题
    doc.add_heading("个人信息登记表", level=1)
    doc.add_heading("基本信息", level=2)

    # 段落
    doc.add_paragraph("请如实填写以下信息。")

    # 表格 1：4 行 4 列，含空格和合并单元格
    t1 = doc.add_table(rows=4, cols=4)
    t1.cell(0, 0).text = "姓名"
    t1.cell(0, 1).text = ""  # 空：可填写
    t1.cell(0, 2).text = "性别"
    t1.cell(0, 3).text = ""  # 空：可填写
    t1.cell(1, 0).text = "出生日期"
    t1.cell(1, 1).text = ""  # 空：可填写
    t1.cell(1, 2).text = "籍贯"
    t1.cell(1, 3).text = ""  # 空：可填写
    t1.cell(2, 0).text = "联系电话"
    t1.cell(2, 1).text = ""  # 空：可填写
    # 第 3 行：把 (3,0) 和 (3,1) 合并作为「备注」标签，(3,2)(3,3) 合并作为空值
    a = t1.cell(3, 0)
    b = t1.cell(3, 1)
    a.merge(b)
    a.text = "备注"
    c = t1.cell(3, 2)
    d = t1.cell(3, 3)
    c.merge(d)  # 合并后的空格：可填写

    # 表格 2：纯展示，无空格
    t2 = doc.add_table(rows=2, cols=2)
    t2.cell(0, 0).text = "项目"
    t2.cell(0, 1).text = "说明"
    t2.cell(1, 0).text = "填表日期"
    t2.cell(1, 1).text = "2026年8月"

    out = tmp_path / "sample.docx"
    doc.save(str(out))
    return out


def test_parse_structure_basic(sample_docx: Path) -> None:
    """验证基本结构：标题、段落、表格、可填写位置。"""
    svc = FormStructureService()
    result = svc.parse_structure(str(sample_docx), filename="sample.docx")

    assert result["format"] == "docx"
    assert result["filename"] == "sample.docx"
    assert result["title"] == "个人信息登记表"


def test_parse_structure_headings(sample_docx: Path) -> None:
    """验证标题提取与层级。"""
    result = FormStructureService().parse_structure(str(sample_docx), filename="sample.docx")
    headings = result["headings"]

    assert len(headings) >= 2
    texts = [h["text"] for h in headings]
    assert "个人信息登记表" in texts
    assert "基本信息" in texts
    # 层级正确
    h1 = next(h for h in headings if h["text"] == "个人信息登记表")
    assert h1["level"] == 1
    h2 = next(h for h in headings if h["text"] == "基本信息")
    assert h2["level"] == 2


def test_parse_structure_paragraphs(sample_docx: Path) -> None:
    """验证段落提取。"""
    result = FormStructureService().parse_structure(str(sample_docx), filename="sample.docx")
    paragraphs = result["paragraphs"]
    texts = [p["text"] for p in paragraphs]
    assert "请如实填写以下信息。" in texts


def test_parse_structure_tables(sample_docx: Path) -> None:
    """验证表格解析：2 个表，第一个表有合并单元格。"""
    result = FormStructureService().parse_structure(str(sample_docx), filename="sample.docx")
    tables = result["tables"]

    assert len(tables) == 2
    # 第一个表 4 行 4 列
    assert tables[0]["rows"] == 4
    assert tables[0]["cols"] == 4
    # 含「姓名」单元格
    cell_texts = [c["text"] for c in tables[0]["cells"]]
    assert "姓名" in cell_texts
    assert "备注" in cell_texts
    # 有合并单元格
    assert len(tables[0]["merged_ranges"]) >= 2  # 备注(横合并) + 空值(横合并)


def test_parse_structure_header_rows(sample_docx: Path) -> None:
    """验证表头行数启发式：form 式表=0（标签与空格同行），展示表=1（首行表头）。"""
    result = FormStructureService().parse_structure(str(sample_docx), filename="sample.docx")
    tables = result["tables"]
    # 表0：form 式（第0行就有空格）→ header_rows=0
    assert tables[0]["header_rows"] == 0
    # 表1：展示表（无空格）→ header_rows=1
    assert tables[1]["header_rows"] == 1


def test_parse_structure_multi_level_header(tmp_path: Path) -> None:
    """验证多级表头：合并跨列表头 + 子表头 + 数据行 → header_rows=2。"""
    doc = Document()
    t = doc.add_table(rows=3, cols=4)
    # 第0行：跨列表头（合并 (0,0)+(0,1) 和 (0,2)+(0,3)）
    a = t.cell(0, 0); b = t.cell(0, 1); a.merge(b); a.text = "个人信息"
    c = t.cell(0, 2); d = t.cell(0, 3); c.merge(d); c.text = "工作信息"
    # 第1行：子表头
    t.cell(1, 0).text = "姓名"
    t.cell(1, 1).text = "性别"
    t.cell(1, 2).text = "公司"
    t.cell(1, 3).text = "职位"
    # 第2行：数据行（空格）
    # （留空即可）
    p = tmp_path / "multi_header.docx"
    doc.save(str(p))
    result = FormStructureService().parse_structure(str(p), filename="multi_header.docx")
    table = result["tables"][0]
    # 前2行是表头（无空格），第2行有空格 → header_rows=2
    assert table["header_rows"] == 2
    # 有多级表头的合并范围
    assert len(table["merged_ranges"]) >= 2


def test_parse_structure_template_fillable(tmp_path: Path) -> None:
    """验证模板填空栏（日期/地址/盖章）被识别为可填写位置。"""
    doc = Document()
    t = doc.add_table(rows=4, cols=2)
    # 日期模板
    t.cell(0, 0).text = "活动时间"
    t.cell(0, 1).text = "年  月  日—年  月  日"
    # 地址模板
    t.cell(1, 0).text = "实践基地"
    t.cell(1, 1).text = "省    市    县（区）   镇    村"
    # 盖章意见栏
    t.cell(2, 0).text = "辅导员意见"
    t.cell(2, 1).text = "盖章：\n年  月  日"
    # 纯空格
    t.cell(3, 0).text = "备注"
    # cell(3,1) 留空
    p = tmp_path / "template.docx"
    doc.save(str(p))
    result = FormStructureService().parse_structure(str(p), filename="template.docx")
    fields = result["fields"]
    # 4 个可填写位置：日期模板 + 地址模板 + 盖章 + 空格
    assert len(fields) == 4
    names = {f["name"] for f in fields}
    assert "活动时间" in names
    assert "实践基地" in names
    assert "辅导员意见" in names
    assert "备注" in names
    # 标签格不应出现在可填写列表
    assert "年  月  日" not in names


def test_parse_structure_multi_row_label_inheritance(tmp_path: Path) -> None:
    """验证多行向上继承字段名：行8空格继承行6表头（跨2行）。"""
    doc = Document()
    t = doc.add_table(rows=3, cols=2)
    # 行0：表头
    t.cell(0, 0).text = "姓名"
    t.cell(0, 1).text = "性别"
    # 行1：空数据行
    # 行2：空数据行（需要跨2行向上继承）
    p = tmp_path / "inherit.docx"
    doc.save(str(p))
    result = FormStructureService().parse_structure(str(p), filename="inherit.docx")
    fields = result["fields"]
    # 行1和行2各2个空格 = 4个字段
    assert len(fields) == 4
    # 行2的字段名应继承行0表头（不是(未知)）
    row2_fields = [f for f in fields if f["locator"]["row"] == 2]
    assert len(row2_fields) == 2
    names = {f["name"] for f in row2_fields}
    assert "姓名" in names
    assert "性别" in names


def test_parse_structure_header_not_fillable(tmp_path: Path) -> None:
    """验证纯标签表头行不被标记为可填写位置。"""
    doc = Document()
    t = doc.add_table(rows=2, cols=3)
    # 行0：纯表头（全部非空）
    t.cell(0, 0).text = "姓名"
    t.cell(0, 1).text = "年龄"
    t.cell(0, 2).text = "电话"
    # 行1：空数据行
    p = tmp_path / "header.docx"
    doc.save(str(p))
    result = FormStructureService().parse_structure(str(p), filename="header.docx")
    fields = result["fields"]
    # 只有行1的3个空格是可填写位，行0表头不算
    assert len(fields) == 3
    for f in fields:
        assert f["locator"]["row"] == 1  # 都是行1
    # 字段名继承行0表头
    names = {f["name"] for f in fields}
    assert names == {"姓名", "年龄", "电话"}


def test_parse_structure_fill_kind(tmp_path: Path) -> None:
    """验证 fill_kind：空格=empty，模板填空栏=template。"""
    doc = Document()
    t = doc.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "活动时间"
    t.cell(0, 1).text = "年  月  日—年  月  日"  # 模板
    t.cell(1, 0).text = "姓名"
    # cell(1,1) 留空 → empty
    t.cell(2, 0).text = "辅导员意见"
    t.cell(2, 1).text = "盖章：\n年  月  日"  # 模板
    p = tmp_path / "fill_kind.docx"
    doc.save(str(p))
    result = FormStructureService().parse_structure(str(p), filename="fill_kind.docx")
    fields = result["fields"]
    assert len(fields) == 3
    by_row = {f["locator"]["row"]: f for f in fields}
    assert by_row[0]["fill_kind"] == "template"  # 日期模板
    assert by_row[1]["fill_kind"] == "empty"      # 纯空格
    assert by_row[2]["fill_kind"] == "template"  # 盖章模板


def test_parse_structure_instruction_text(tmp_path: Path) -> None:
    """验证填写说明文字区被识别为 template（签字/长说明/填空提示）。"""
    doc = Document()
    t = doc.add_table(rows=5, cols=2)
    # 行0：签字区
    t.cell(0, 0).text = "年级推荐意见"
    t.cell(0, 1).text = "辅导员签字："
    # 行1：长说明文字
    t.cell(1, 0).text = "曾获奖项"
    t.cell(1, 1).text = "大学期间获得的校级及以上奖励，如果无校级及以上，则院级也可以"
    # 行2：填空提示
    t.cell(2, 0).text = "学生工作经历"
    t.cell(2, 1).text = "时间段＋具体岗位"
    # 行3：纯标签（不应被识别）
    t.cell(3, 0).text = "姓名"
    t.cell(3, 1).text = ""  # empty → 可填写
    # 行4：参考列表（不应被识别为可填写）
    t.cell(4, 0).text = "意向部门"
    t.cell(4, 1).text = "学习部 调研部 权益部 活动部 （至少选择两个）"
    p = tmp_path / "instruction.docx"
    doc.save(str(p))
    result = FormStructureService().parse_structure(str(p), filename="instruction.docx")
    fields = result["fields"]
    # 行0-2 的说明文字区 + 行3 的空格 = 4 个
    assert len(fields) == 4
    by_row = {f["locator"]["row"]: f for f in fields}
    assert by_row[0]["fill_kind"] == "template"  # 签字
    assert by_row[0]["name"] == "年级推荐意见"
    assert by_row[1]["fill_kind"] == "template"  # 长说明
    assert by_row[1]["name"] == "曾获奖项"
    assert by_row[2]["fill_kind"] == "template"  # 填空提示
    assert by_row[2]["name"] == "学生工作经历"
    assert by_row[3]["fill_kind"] == "empty"     # 纯空格
    assert by_row[3]["name"] == "姓名"
    # 行4 参考列表不是可填写位
    assert 4 not in by_row


def test_parse_structure_multi_col_header_inheritance(tmp_path: Path) -> None:
    """验证多级表头字段名继承：纵向合并小节标题不干扰列头继承 + 横合并表头溯源。"""
    doc = Document()
    t = doc.add_table(rows=4, cols=5)
    # 行0：纵向合并的小节标题「家庭成员情况」（跨行0-3，覆盖全部行）
    t.cell(0, 0).text = "家庭成员情况"
    t.cell(0, 0).merge(t.cell(3, 0))
    # 行0：列头（col1-4），其中 col3-4 横合并为「年收入」
    t.cell(0, 1).text = "姓名"
    t.cell(0, 2).text = "年龄"
    t.cell(0, 3).text = "年收入"
    t.cell(0, 3).merge(t.cell(0, 4))
    # 行1-3：数据行（空格）
    p = tmp_path / "multi_col_header.docx"
    doc.save(str(p))
    result = FormStructureService().parse_structure(str(p), filename="multi_col_header.docx")
    fields = result["fields"]
    by_pos = {(f["locator"]["row"], f["locator"]["col"]): f for f in fields}
    # 行1 col1 应继承「姓名」，不是「家庭成员情况」
    assert by_pos[(1, 1)]["name"] == "姓名"
    assert by_pos[(1, 2)]["name"] == "年龄"
    # 行1 col3 应继承横合并表头「年收入」
    assert by_pos[(1, 3)]["name"] == "年收入"
    # 行1 col4 是横合并延续格，也应继承「年收入」
    assert by_pos[(1, 4)]["name"] == "年收入"
    # 不应出现「家庭成员情况」作为字段名
    for f in fields:
        assert f["name"] != "家庭成员情况"


def test_parse_structure_fields_and_locator(sample_docx: Path) -> None:
    """验证可填写位置识别 + locator 契约一致性。"""
    result = FormStructureService().parse_structure(str(sample_docx), filename="sample.docx")
    fields = result["fields"]

    # 至少识别出几个空格
    assert len(fields) >= 5

    # 每个 field 的 locator 必须符合 WordCellLocator 契约
    for f in fields:
        assert f["locator"]["kind"] == "word_cell"
        # 用契约模型校验（不抛异常即通过）
        loc = WordCellLocator(**f["locator"])
        assert loc.table_index >= 0
        assert loc.row >= 0
        assert loc.col >= 0

    # 「姓名」右侧的空格应能推断出字段名「姓名」
    name_field = next(
        (f for f in fields if f["locator"]["table_index"] == 0
         and f["locator"]["row"] == 0 and f["locator"]["col"] == 1),
        None,
    )
    assert name_field is not None
    assert name_field["name"] == "姓名"


def test_parse_structure_field_ids_unique(sample_docx: Path) -> None:
    """每个可填写位置的 id 应唯一。"""
    result = FormStructureService().parse_structure(str(sample_docx), filename="sample.docx")
    ids = [f["id"] for f in result["fields"]]
    assert len(ids) == len(set(ids))


def test_parse_structure_structure_notes(sample_docx: Path) -> None:
    """验证结构备注。"""
    result = FormStructureService().parse_structure(str(sample_docx), filename="sample.docx")
    notes = result["structure_notes"]
    assert any("表格" in n for n in notes)


def test_parse_structure_result_serializable(sample_docx: Path) -> None:
    """结果应能序列化为 JSON（供 API 返回）。"""
    result = FormStructureService().parse_structure(str(sample_docx), filename="sample.docx")
    # 不抛异常即通过
    json.dumps(result, ensure_ascii=False)


def test_parse_structure_file_not_found() -> None:
    """文件不存在应抛 FileNotFoundError。"""
    svc = FormStructureService()
    with pytest.raises(FileNotFoundError):
        svc.parse_structure("nonexistent.docx", filename="nonexistent.docx")


def test_parse_structure_unsupported_format(tmp_path: Path) -> None:
    """非 Word 文件应抛 ValueError。"""
    f = tmp_path / "not_a_doc.txt"
    f.write_text("hello", encoding="utf-8")
    svc = FormStructureService()
    with pytest.raises(ValueError):
        svc.parse_structure(str(f), filename="not_a_doc.txt")
